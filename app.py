import os
import json
import hashlib
import time
from dotenv import load_dotenv
from document_processor import process_document_folder
from langchain.docstore.document import Document
from langchain_community.vectorstores import FAISS as LC_FAISS
from langchain_openai import ChatOpenAI
from langchain.chains.conversational_retrieval.base import ConversationalRetrievalChain
import gradio as gr
from langchain.embeddings.base import Embeddings
from document_manager import DocumentManager
from upload_ui import initialize_document_manager, create_upload_tab
import cv2
import pytesseract
import numpy as np
from PIL import Image
from docx import Document
import fitz  # PyMuPDF

# Load environment variables from a .env file
load_dotenv()

# Retrieve configuration from environment variables
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
DATA_FOLDER = os.getenv("DATA_FOLDER", "data")
INDEXES_FOLDER = os.getenv("INDEXES_FOLDER", "indexes")

# Путь к Tesseract (измените на актуальный путь на вашей системе)
pytesseract.pytesseract.tesseract_cmd = r"E:/Games/tesseract.exe"

# Папка для сохранения отладочных изображений
DEBUG_FOLDER = "debug_stamps"
os.makedirs(DEBUG_FOLDER, exist_ok=True)

# Глобальные переменные для отчёта по проверке документов
detailed_missing_report = {}
good_documents = []


# --------------------
# Custom Embeddings Class
# --------------------
class MyEmbeddings(Embeddings):
    def __init__(self):
        from sentence_transformers import SentenceTransformer
        self.model = SentenceTransformer('all-MiniLM-L12-v2')

    def embed_documents(self, texts):
        return self.model.encode(texts, convert_to_numpy=True).tolist()

    def embed_query(self, text):
        return self.model.encode([text], convert_to_numpy=True)[0].tolist()


embeddings = MyEmbeddings()


# --------------------
# Document Verification Functions
# --------------------
def extract_images_from_pdf(pdf_path):
    """Извлекает изображения из PDF."""
    images = []
    try:
        doc = fitz.open(pdf_path)
        for page in doc:
            pix = page.get_pixmap(dpi=300)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            images.append(cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR))
    except Exception as e:
        print(f"❌ Ошибка при открытии PDF: {e}")
    return images


def extract_images_from_docx(docx_path):
    """Извлекает изображение из DOCX, рендеря текст."""
    doc = Document(docx_path)
    full_text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
    img = np.ones((2000, 1600), dtype=np.uint8) * 255
    y0 = 50
    for i, line in enumerate(full_text.split('\n')):
        y = y0 + i * 30
        if y >= img.shape[0] - 50:
            break
        cv2.putText(img, line, (50, y), cv2.FONT_HERSHEY_SIMPLEX, 0.8, (0,), 2)
    img_bgr = cv2.cvtColor(img, cv2.COLOR_GRAY2BGR)
    return [img_bgr]


def filter_logo(contours, blue_mask, gray):
    """Фильтрует контуры, чтобы отличить логотип университета от печати."""
    filtered_contours = []
    logo_regions = []

    for cnt in contours:
        if cv2.contourArea(cnt) < 400:
            continue

        x, y, w, h = cv2.boundingRect(cnt)
        x_extended = max(0, x - 20)
        y_extended = max(0, y - 20)
        w_extended = min(gray.shape[1] - x_extended, w + 40)
        h_extended = min(gray.shape[0] - y_extended, h + 40)

        roi = gray[y_extended:y_extended + h_extended, x_extended:x_extended + w_extended]
        if roi.size == 0:
            continue

        text = pytesseract.image_to_string(roi, lang='eng').lower()
        logo_keywords = ['astana', 'university', 'университет']
        if any(kw in text for kw in logo_keywords):
            logo_regions.append((x_extended, y_extended, w_extended, h_extended))
        else:
            (_, _), radius = cv2.minEnclosingCircle(cnt)
            if radius > 30:
                mask = np.zeros_like(gray)
                cv2.drawContours(mask, [cnt], 0, 255, -1)
                blue_pixels = cv2.bitwise_and(blue_mask, mask)
                blue_ratio = np.count_nonzero(blue_pixels) / np.count_nonzero(mask) if np.count_nonzero(mask) > 0 else 0
                if blue_ratio > 0.1:
                    filtered_contours.append(cnt)

    for cnt in contours:
        if cv2.contourArea(cnt) < 400:
            continue

        x, y, w, h = cv2.boundingRect(cnt)
        is_overlapping = False

        for l_x, l_y, l_w, l_h in logo_regions:
            if (x < l_x + l_w and x + w > l_x and
                    y < l_y + l_h and y + h > l_y):
                is_overlapping = True
                break

        if not is_overlapping:
            (_, _), radius = cv2.minEnclosingCircle(cnt)
            perimeter = cv2.arcLength(cnt, True)
            approx = cv2.approxPolyDP(cnt, 0.04 * perimeter, True)

            if (len(approx) > 6 and radius > 20) or radius > 40:
                mask = np.zeros_like(gray)
                cv2.drawContours(mask, [cnt], 0, 255, -1)
                blue_pixels = cv2.bitwise_and(blue_mask, mask)
                blue_ratio = np.count_nonzero(blue_pixels) / np.count_nonzero(mask) if np.count_nonzero(mask) > 0 else 0
                if blue_ratio > 0.1:
                    filtered_contours.append(cnt)

    return filtered_contours


def detect_signature(image):
    """Обнаруживает подпись на изображении через контуры."""
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    thresh = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                   cv2.THRESH_BINARY_INV, 11, 2)
    kernel = np.ones((2, 2), np.uint8)
    opening = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, kernel, iterations=1)

    contours, _ = cv2.findContours(opening, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    signature_contours = []
    for cnt in contours:
        area = cv2.contourArea(cnt)
        if 50 < area < 1000:
            perimeter = cv2.arcLength(cnt, True)
            approx = cv2.approxPolyDP(cnt, 0.01 * perimeter, True)
            if len(approx) > 4:
                signature_contours.append(cnt)

    return len(signature_contours) > 5


def detect_concentric_patterns(mask, min_rings=3):
    """Обнаруживает концентрические узоры, характерные для печатей."""
    dist = cv2.distanceTransform(mask, cv2.DIST_L2, 5)
    cv2.normalize(dist, dist, 0, 1.0, cv2.NORM_MINMAX)
    hist = cv2.calcHist([dist * 255], [0], None, [50], [1, 255])

    peaks = []
    for i in range(1, len(hist) - 1):
        if hist[i] > hist[i - 1] and hist[i] > hist[i + 1] and hist[i] > 5:
            peaks.append(i)

    return len(peaks) >= min_rings


def detect_stamp(image):
    """Обнаруживает печать в правой верхней части изображения."""
    debug_original = image.copy()
    h, w = image.shape[:2]
    right_boundary = int(w * 0.6)
    right_part = image[:, right_boundary:]

    debug_right_part = right_part.copy()
    hsv = cv2.cvtColor(right_part, cv2.COLOR_BGR2HSV)

    lower_blue = np.array([90, 50, 50])
    upper_blue = np.array([130, 255, 255])
    blue_mask = cv2.inRange(hsv, lower_blue, upper_blue)

    gray = cv2.cvtColor(right_part, cv2.COLOR_BGR2GRAY)
    thresh = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                   cv2.THRESH_BINARY_INV, 11, 2)

    combined_mask = cv2.bitwise_or(thresh, blue_mask)
    kernel = np.ones((3, 3), np.uint8)
    opening = cv2.morphologyEx(combined_mask, cv2.MORPH_OPEN, kernel, iterations=1)
    closing = cv2.morphologyEx(opening, cv2.MORPH_CLOSE, kernel, iterations=2)

    contours, _ = cv2.findContours(closing, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    debug_contours = cv2.cvtColor(opening, cv2.COLOR_GRAY2BGR)
    cv2.drawContours(debug_contours, contours, -1, (0, 255, 0), 2)

    stamp_contours = filter_logo(contours, blue_mask, gray)
    for cnt in stamp_contours:
        cv2.drawContours(debug_contours, [cnt], 0, (0, 0, 255), 2)

    debug_blue = cv2.cvtColor(blue_mask, cv2.COLOR_GRAY2BGR)
    has_signature = detect_signature(right_part)

    text = pytesseract.image_to_string(right_part, lang='rus+eng').lower()
    stamp_keywords = ['печать', 'м.п.', 'seal', 'stamp', 'декан', 'подпись', 'dean', 'signature']
    text_has_stamp_indicators = any(kw in text for kw in stamp_keywords)
    approved_text = 'approved' in text or 'утверждено' in text

    definite_stamp_contours = []
    for cnt in stamp_contours:
        perimeter = cv2.arcLength(cnt, True)
        area = cv2.contourArea(cnt)
        if area < 500:
            continue

        circularity = 4 * np.pi * area / (perimeter * perimeter) if perimeter > 0 else 0
        (_, _), radius = cv2.minEnclosingCircle(cnt)

        mask = np.zeros_like(gray)
        cv2.drawContours(mask, [cnt], 0, 255, -1)
        blue_pixels = cv2.bitwise_and(blue_mask, mask)
        blue_ratio = np.count_nonzero(blue_pixels) / np.count_nonzero(mask) if np.count_nonzero(mask) > 0 else 0

        has_concentric = detect_concentric_patterns(mask)
        if (circularity > 0.7 and blue_ratio > 0.15 and radius > 30) or (has_concentric and blue_ratio > 0.1):
            definite_stamp_contours.append(cnt)

    stamp_like_objects = []
    for cnt in contours:
        if cv2.contourArea(cnt) < 400:
            continue

        perimeter = cv2.arcLength(cnt, True)
        area = cv2.contourArea(cnt)
        circularity = 4 * np.pi * area / (perimeter * perimeter) if perimeter > 0 else 0

        mask = np.zeros_like(gray)
        cv2.drawContours(mask, [cnt], 0, 255, -1)
        has_concentric_pattern = detect_concentric_patterns(mask)

        blue_pixels = cv2.bitwise_and(blue_mask, mask)
        blue_ratio = np.count_nonzero(blue_pixels) / np.count_nonzero(mask) if np.count_nonzero(mask) > 0 else 0
        (_, _), radius = cv2.minEnclosingCircle(cnt)

        if (circularity > 0.6 and blue_ratio > 0.1 and radius > 25) or (has_concentric_pattern and blue_ratio > 0.05):
            stamp_like_objects.append(cnt)

    has_definite_stamp = len(definite_stamp_contours) > 0
    has_probable_stamp = (len(stamp_like_objects) > 0 and (has_signature or text_has_stamp_indicators))
    approval_section_exists = approved_text

    has_stamp = has_definite_stamp or (has_probable_stamp and approval_section_exists)

    debug_info = {
        "original": debug_original,
        "right_part": debug_right_part,
        "threshold": thresh,
        "blue_mask": debug_blue,
        "contours": debug_contours
    }

    return has_stamp, has_signature, debug_info


def analyze_document(file_path):
    """Анализирует документ на наличие подписи и печати."""
    global detailed_missing_report, good_documents
    file_name = os.path.basename(file_path)
    print(f"\n📄 Обработка файла: {file_name}")

    try:
        if file_path.endswith('.pdf'):
            images = extract_images_from_pdf(file_path)
        elif file_path.endswith('.docx'):
            images = extract_images_from_docx(file_path)
        else:
            return f"⛔ {file_name}: Неподдерживаемый формат."

        if not images or len(images) == 0:
            return f"⚠️ {file_name}: Не удалось извлечь страницы."

        first_page = images[0]
        h, w = first_page.shape[:2]
        upper_part = first_page[0:int(h * 0.4), :]

        has_stamp, has_signature, debug_images = detect_stamp(upper_part)

        base_name = os.path.splitext(file_name)[0]
        cv2.imwrite(os.path.join(DEBUG_FOLDER, f"{base_name}_original.png"), debug_images["original"])
        cv2.imwrite(os.path.join(DEBUG_FOLDER, f"{base_name}_right_part.png"), debug_images["right_part"])
        cv2.imwrite(os.path.join(DEBUG_FOLDER, f"{base_name}_threshold.png"), debug_images["threshold"])
        cv2.imwrite(os.path.join(DEBUG_FOLDER, f"{base_name}_blue_mask.png"), debug_images["blue_mask"])
        cv2.imwrite(os.path.join(DEBUG_FOLDER, f"{base_name}_contours.png"), debug_images["contours"])

        issues = []
        if not has_stamp:
            issues.append("Печать на первой странице в правой части")
        if not has_signature:
            issues.append("Подпись на первой странице в правой части")

        if issues:
            detailed_missing_report[file_name] = issues
            return f"❌ {file_name}:\n" + "\n".join(f" - {issue}" for issue in issues)
        else:
            good_documents.append(file_name)
            return f"✅ {file_name}: Печать и подпись найдены."

    except Exception as e:
        return f"❌ {file_name}: Ошибка при обработке: {e}"


def verify_documents(folder_path):
    """Обрабатывает все PDF и DOCX файлы в указанной папке и возвращает результаты проверки."""
    global detailed_missing_report, good_documents
    detailed_missing_report = {}
    good_documents = []
    results = []

    # Проверяем, существует ли папка
    if not os.path.isdir(folder_path):
        return "❌ Указанный путь не является папкой или не существует."

    # Получаем список всех PDF и DOCX файлов в папке
    files = [os.path.join(folder_path, f) for f in os.listdir(folder_path)
             if f.lower().endswith(('.pdf', '.docx')) and not f.startswith('~$')]

    if not files:
        return "⚠️ В папке нет PDF или DOCX файлов."

    for file_path in files:
        result = analyze_document(file_path)
        results.append(result)

    # Формируем финальный отчёт
    final_report = "## Результаты проверки\n\n"
    if detailed_missing_report:
        final_report += "### Документы с проблемами:\n"
        for doc, issues in detailed_missing_report.items():
            final_report += f"- **{doc}**:\n" + "\n".join(f"  - {issue}" for issue in issues) + "\n"
    else:
        final_report += "✅ Во всех документах печать и подпись присутствуют.\n"

    if good_documents:
        final_report += "\n### Документы без проблем:\n"
        for doc in good_documents:
            final_report += f"- **{doc}**\n"

    return final_report


# --------------------
# File Modification Check Functions
# --------------------
def load_file_modification_info(data_folder):
    mod_info = {}
    for root, _, files in os.walk(data_folder):
        for file in files:
            if file.lower().endswith(('.docx', '.pdf', '.txt')):
                path = os.path.join(root, file)
                mod_info[path] = os.path.getmtime(path)
    return mod_info


def load_previous_mod_info(mod_file_path):
    if os.path.exists(mod_file_path):
        with open(mod_file_path, 'r') as f:
            return json.load(f)
    return None


def save_mod_info(mod_info, mod_file_path):
    with open(mod_file_path, 'w') as f:
        json.dump(mod_info, f)


# --------------------
# Index Building / Rebuilding Functions
# --------------------
def rebuild_index(data_folder, indexes_folder):
    print("[INFO] Rebuilding index from documents...")
    os.makedirs(data_folder, exist_ok=True)
    os.makedirs(indexes_folder, exist_ok=True)

    chunks = process_document_folder(
        data_folder,
        min_words_per_page=100,
        target_chunk_size=512,
        min_chunk_size=256,
        overlap_size=150
    )

    if not chunks:
        print("[WARNING] No document chunks were generated. Check your data folder.")
        empty_doc = Document()
        empty_doc.page_content = "Empty index"
        empty_doc.metadata = {}
        empty_doc.id = "empty_index_0"
        vectorstore = LC_FAISS.from_documents([empty_doc], embeddings)
        vectorstore.save_local(indexes_folder)
        return vectorstore

    chunks_json = os.path.join(indexes_folder, "updated_chunks_1.json")
    with open(chunks_json, 'w', encoding='utf-8') as f:
        json.dump(chunks, f, ensure_ascii=False, indent=2)
    print("[INFO] Saved chunks to updated_chunks_1.json")

    # Создаём объекты Document с уникальным id
    docs = []
    for i, ch in enumerate(chunks):
        doc = Document()
        doc.page_content = ch["text"]
        doc.metadata = ch["metadata"]
        doc.id = hashlib.md5(f"{ch['text']}_{i}".encode()).hexdigest()
        docs.append(doc)

    vectorstore = LC_FAISS.from_documents(docs, embeddings)
    vectorstore.save_local(indexes_folder)

    total_chunks = len(chunks)
    unique_docs = len({ch["metadata"]["file_path"] for ch in chunks})
    total_tokens = sum(ch["metadata"].get("token_count", 0) for ch in chunks)
    avg_tokens = total_tokens / total_chunks if total_chunks > 0 else 0
    print(f"[INFO] Created {total_chunks} chunks from {unique_docs} documents")
    print(f"[INFO] Average chunk token count: {avg_tokens:.2f}")
    print(f"[INFO] Total number of tokens in all chunks: {total_tokens}")
    return vectorstore


def load_document_metadata_hash(metadata_file_path):
    if os.path.exists(metadata_file_path):
        with open(metadata_file_path, 'rb') as f:
            content = f.read()
        return hashlib.md5(content).hexdigest()
    return None


def load_or_rebuild_vectorstore(data_folder, indexes_folder):
    fingerprint_file = os.path.join(indexes_folder, "index_fingerprint.json")

    current_fingerprint = {}
    for root, _, files in os.walk(data_folder):
        for file in files:
            if file.lower().endswith(('.docx', '.pdf', '.txt')):
                path = os.path.join(root, file)
                current_fingerprint[path] = os.path.getmtime(path)

    fingerprint_hash = hashlib.md5(json.dumps(current_fingerprint, sort_keys=True).encode()).hexdigest()

    previous_fingerprint = None
    if os.path.exists(fingerprint_file):
        try:
            with open(fingerprint_file, 'r') as f:
                previous_fingerprint = json.load(f)
        except (json.JSONDecodeError, IOError) as e:
            print(f"[WARNING] Could not read fingerprint file: {e}")

    index_path = os.path.join(indexes_folder, "index.faiss")
    metadata_path = os.path.join(indexes_folder, "metadata.json")

    print(f"[DEBUG] Current fingerprint: {fingerprint_hash}")
    print(f"[DEBUG] Previous fingerprint: {previous_fingerprint}")

    if (os.path.exists(index_path) and
            os.path.exists(metadata_path) and
            previous_fingerprint == fingerprint_hash):
        try:
            vectorstore = LC_FAISS.load_local(indexes_folder, embeddings)
            print("[INFO] Loaded existing vectorstore.")
            return vectorstore
        except Exception as e:
            print(f"[INFO] Failed to load existing vectorstore: {e}")

    print("[INFO] Building/rebuilding index from documents...")
    vectorstore = rebuild_index(data_folder, indexes_folder)

    with open(fingerprint_file, 'w') as f:
        json.dump(fingerprint_hash, f)

    return vectorstore


# --------------------
# Custom Prompt for QA
# --------------------
def get_qa_prompt_template():
    return """You are a helpful university assistant. Use the following pieces of retrieved context to answer the question. 
    If you don't know the answer, just say you don't know. DO NOT try to make up an answer. Write full answer in 2-3 paragraphs.

    {context}

    Question: {question}

    Answer the question clearly and helpfully. DO NOT include any sources, citations, or references at the end of your answer. Write full answer in 2-3 paragraphs.
    """


# --------------------
# Build Conversational Retrieval Chain
# --------------------
llm = ChatOpenAI(openai_api_key=OPENAI_API_KEY, temperature=0)
if not os.path.exists(INDEXES_FOLDER):
    os.makedirs(INDEXES_FOLDER)
vectorstore = load_or_rebuild_vectorstore(DATA_FOLDER, INDEXES_FOLDER)

from langchain.prompts.prompt import PromptTemplate
from langchain.chains.question_answering import load_qa_chain

qa_prompt = PromptTemplate(
    template=get_qa_prompt_template(),
    input_variables=["context", "question"]
)

qa_chain = ConversationalRetrievalChain.from_llm(
    llm,
    retriever=vectorstore.as_retriever(),
    return_source_documents=True,
    combine_docs_chain_kwargs={"prompt": qa_prompt}
)


# --------------------
# Chat State Management Class
# --------------------
class ChatAssistant:
    def __init__(self, qa_chain):
        self.chat_history = []
        self.qa = qa_chain

    def _to_chain_history(self):
        pairs = []
        user_text = None
        for msg in self.chat_history:
            if msg["role"] == "user":
                user_text = msg["content"]
            elif msg["role"] == "assistant":
                if user_text is not None:
                    answer = msg["content"]
                    if "Sources:" in answer:
                        answer = answer.split("Sources:")[0].strip()
                    pairs.append((user_text, answer))
                    user_text = None
                else:
                    pairs.append(("", msg["content"]))
        return pairs

    def _extract_sources(self, source_docs):
        sources = {}
        for doc in source_docs:
            metadata = doc.metadata
            if "file_name" not in metadata:
                continue

            file_name = metadata.get("file_name")

            if metadata.get("file_type") == "pdf" and "page_number" in metadata:
                if file_name not in sources:
                    sources[file_name] = []
                if metadata["page_number"] not in sources[file_name]:
                    sources[file_name].append(metadata["page_number"])
            else:
                if file_name not in sources:
                    sources[file_name] = []

        formatted_sources = []
        for file_name, pages in sources.items():
            if pages:
                pages.sort()
                page_ranges = []
                start = pages[0]
                end = pages[0]

                for i in range(1, len(pages)):
                    if pages[i] == end + 1:
                        end = pages[i]
                    else:
                        if start == end:
                            page_ranges.append(str(start))
                        else:
                            page_ranges.append(f"{start}-{end}")
                        start = end = pages[i]

                if start == end:
                    page_ranges.append(str(start))
                else:
                    page_ranges.append(f"{start}-{end}")

                formatted_sources.append(f"- {file_name} (Page {', '.join(page_ranges)})")
            else:
                formatted_sources.append(f"- {file_name}")

        return formatted_sources

    def convchain(self, query):
        if not query:
            return self.chat_history
        if not self.chat_history or self.chat_history[-1]["role"] != "user" or self.chat_history[-1][
            "content"] != query:
            chain_history = self._to_chain_history()
            result = self.qa({"question": query, "chat_history": chain_history})
            answer = result.get("answer", "")

            source_docs = result.get("source_documents", [])
            if source_docs:
                sources = self._extract_sources(source_docs)
                if sources:
                    sources_text = "\n\nSources:\n" + "\n".join(sources)
                    answer += sources_text

            self.chat_history.append({"role": "user", "content": query})
            self.chat_history.append({"role": "assistant", "content": answer})
        return self.chat_history

    def clr_history(self):
        self.chat_history = []
        return []

    def get_history_text(self):
        if not self.chat_history:
            return "No conversation history yet."

        formatted_history = []
        for msg in self.chat_history:
            role = "👤 User" if msg["role"] == "user" else "🤖 Assistant"
            content = msg["content"]
            if msg["role"] == "assistant" and "Sources:" in content:
                content = content.split("Sources:")[0].strip()

            formatted_history.append(f"**{role}**:\n{content}")

        return "\n\n".join(formatted_history)


# Instantiate our assistant
assistant = ChatAssistant(qa_chain)
document_manager = initialize_document_manager(DATA_FOLDER)


# --------------------
# Gradio Callback Functions
# --------------------
def process_query(query):
    updated_messages = assistant.convchain(query)
    return updated_messages, ""


def clear_history_callback():
    assistant.clr_history()
    return [], "No conversation history yet."


# --------------------
# Gradio Interface
# --------------------
with gr.Blocks(css="""
    .history-box {
        border: 1px solid #333;
        border-radius: 10px;
        background-color: #2d2d2d;
        padding: 15px;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
        color: #f0f0f0;
    }
    .history-box p {
        margin-bottom: 10px;
    }
    .history-box strong {
        color: #ffffff;
    }
""") as demo:
    gr.Markdown("# University Chat Assistant")
    with gr.Tabs():
        with gr.TabItem("Conversation"):
            chatbot = gr.Chatbot(label="Conversation", type="messages", height=500)
            with gr.Row():
                query_input = gr.Textbox(placeholder="Type your message...", show_label=False, container=False)
            query_input.submit(fn=process_query, inputs=query_input, outputs=[chatbot, query_input])

        with gr.TabItem("Chat History"):
            history_text = gr.Markdown(value="No conversation history yet.", elem_classes=["history-box"])
            with gr.Row():
                refresh_history_btn = gr.Button("Refresh History")
                clear_history_btn = gr.Button("Clear History")

            refresh_history_btn.click(fn=lambda: assistant.get_history_text(), outputs=history_text)
            clear_history_btn.click(fn=lambda: (assistant.clr_history(), "No conversation history yet."),
                                    outputs=history_text)

        with gr.TabItem("Upload"):
            document_list = create_upload_tab(document_manager, embeddings, vectorstore,
                                              lambda: load_or_rebuild_vectorstore(DATA_FOLDER, INDEXES_FOLDER))

        with gr.TabItem("Document Verification"):
            gr.Markdown(
                "### Enter the path to a folder containing PDF or DOCX files to check for stamps and signatures")
            folder_input = gr.Textbox(label="Folder Path", placeholder="e.g., C:/Users/Админ/Desktop/Syllabus")
            verify_button = gr.Button("Verify Folder")
            result_output = gr.Markdown()
            verify_button.click(fn=verify_documents, inputs=folder_input, outputs=result_output)

if __name__ == "__main__":
    demo.launch()
